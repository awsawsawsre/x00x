from pathlib import Path
from datetime import datetime
import re

import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog

from PIL import Image
import piexif
from pypdf import PdfReader, PdfWriter
from docx import Document


class MetadataGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Metadata Viewer / Editor (Images, PDF, Word)")

        self.current_path = None
        self.current_meta = {}
        self.exif_dict = None
        self.file_type = "other"  # "image", "pdf", "word", "other"
        self.pending_changes = {}

        self.build_ui()

    def build_ui(self):
        top_frame = tk.Frame(self.root)
        top_frame.pack(fill="x", padx=10, pady=8)

        self.file_label = tk.Label(top_frame, text="No file selected", anchor="w")
        self.file_label.pack(side="left", fill="x", expand=True)

        btn_frame = tk.Frame(self.root)
        btn_frame.pack(fill="x", padx=10, pady=4)

        tk.Button(btn_frame, text="Open file", command=self.load_file).pack(side="left", padx=4)
        tk.Button(btn_frame, text="Strip metadata", command=self.strip_metadata).pack(side="left", padx=4)
        tk.Button(btn_frame, text="Edit metadata", command=self.edit_metadata).pack(side="left", padx=4)
        tk.Button(btn_frame, text="Save edits", command=self.save_edits).pack(side="left", padx=4)

        text_frame = tk.Frame(self.root)
        text_frame.pack(fill="both", expand=True, padx=10, pady=8)

        self.text = tk.Text(text_frame, wrap="none", font=("Consolas", 10))
        self.text.pack(side="left", fill="both", expand=True)

        scroll_y = tk.Scrollbar(text_frame, orient="vertical", command=self.text.yview)
        scroll_y.pack(side="right", fill="y")
        self.text.configure(yscrollcommand=scroll_y.set)

        scroll_x = tk.Scrollbar(self.root, orient="horizontal", command=self.text.xview)
        scroll_x.pack(fill="x", padx=10, pady=(0, 8))
        self.text.configure(xscrollcommand=scroll_x.set)

    def _exif_gps_to_decimal(self, gps_tuple, ref=None):
        """
        gps_tuple like ((deg_num, deg_den), (min_num, min_den), (sec_num, sec_den))
        ref is 'N', 'S', 'E', or 'W'
        """
        (deg_num, deg_den), (min_num, min_den), (sec_num, sec_den) = gps_tuple

        degrees = deg_num / deg_den
        minutes = min_num / min_den
        seconds = sec_num / sec_den

        decimal = degrees + minutes / 60.0 + seconds / 3600.0

        if ref in ("S", "W"):
            decimal = -decimal

        return decimal

    def load_file(self):
        path = filedialog.askopenfilename(title="Choose a file")
        if not path:
            return

        self.current_path = path
        self.file_label.config(text=path)
        self.pending_changes.clear()

        self.current_meta, self.exif_dict, self.file_type = self.read_metadata(path)
        self.render_meta()

    def read_metadata(self, path):
        meta = {}
        exif_dict = None
        file_type = "other"

        p = Path(path)
        suffix = p.suffix.lower()

        # File system info
        try:
            stat = p.stat()
            meta["FS.File name"] = p.name
            meta["FS.Directory"] = str(p.parent)
            meta["FS.Size (bytes)"] = str(stat.st_size)
            meta["FS.Mode (octal)"] = oct(stat.st_mode)
            meta["FS.Device"] = str(stat.st_dev)
            meta["FS.Inode"] = str(getattr(stat, "st_ino", "n/a"))
            meta["FS.Nlink"] = str(getattr(stat, "st_nlink", "n/a"))
            meta["FS.UID"] = str(getattr(stat, "st_uid", "n/a"))
            meta["FS.GID"] = str(getattr(stat, "st_gid", "n/a"))
            meta["FS.Accessed"] = datetime.fromtimestamp(stat.st_atime).isoformat(" ")
            meta["FS.Modified"] = datetime.fromtimestamp(stat.st_mtime).isoformat(" ")
            meta["FS.Created"] = datetime.fromtimestamp(stat.st_ctime).isoformat(" ")
        except Exception as e:
            meta["FS.Error"] = str(e)

        # Images and EXIF
        image_exts = {".jpg", ".jpeg", ".tif", ".tiff", ".png", ".webp", ".heic", ".heif"}
        if suffix in image_exts:
            file_type = "image"
            try:
                img = Image.open(path)
                meta["IMG.Format"] = str(img.format)
                meta["IMG.Mode"] = str(img.mode)
                meta["IMG.Size"] = f"{img.width}x{img.height}"

                for k, v in img.info.items():
                    if k == "exif":
                        continue
                    meta[f"IMG.info.{k}"] = repr(v)

                if "exif" in img.info:
                    exif_dict = piexif.load(img.info["exif"])
                else:
                    try:
                        exif_dict = piexif.load(path)
                    except Exception:
                        exif_dict = None
            except Exception as e:
                meta["IMG.Error"] = str(e)
                exif_dict = None

            if exif_dict:
                if exif_dict.get("thumbnail"):
                    meta["EXIF.thumbnail.length"] = str(len(exif_dict["thumbnail"]))

                # Raw EXIF fields
                for ifd_name, ifd_dict in exif_dict.items():
                    if ifd_name == "thumbnail" or not isinstance(ifd_dict, dict):
                        continue
                    for tag_id, value in ifd_dict.items():
                        tag_info = piexif.TAGS.get(ifd_name, {}).get(tag_id, {})
                        tag_name = tag_info.get("name", f"Tag{tag_id}")
                        key = f"EXIF.{ifd_name}.{tag_name}"

                        if isinstance(value, bytes):
                            try:
                                v_str = value.decode("utf-8", errors="replace")
                            except Exception:
                                v_str = repr(value)
                        else:
                            v_str = str(value)

                        meta[key] = v_str

                # GPS decoded to decimal degrees
                try:
                    gps_ifd = exif_dict.get("GPS")
                    if gps_ifd:
                        lat = gps_ifd.get(piexif.GPSIFD.GPSLatitude)
                        lat_ref = gps_ifd.get(piexif.GPSIFD.GPSLatitudeRef)
                        lon = gps_ifd.get(piexif.GPSIFD.GPSLongitude)
                        lon_ref = gps_ifd.get(piexif.GPSIFD.GPSLongitudeRef)

                        if lat and lon and lat_ref and lon_ref:
                            if isinstance(lat_ref, bytes):
                                lat_ref = lat_ref.decode(errors="ignore")
                            if isinstance(lon_ref, bytes):
                                lon_ref = lon_ref.decode(errors="ignore")

                            lat_dec = self._exif_gps_to_decimal(lat, lat_ref)
                            lon_dec = self._exif_gps_to_decimal(lon, lon_ref)

                            meta["EXIF.GPS.decimal_latitude"] = f"{lat_dec:.6f}"
                            meta["EXIF.GPS.decimal_longitude"] = f"{lon_dec:.6f}"
                except Exception as e:
                    meta["EXIF.GPS.error"] = f"Failed to parse GPS: {e}"

        # PDF
        elif suffix == ".pdf":
            file_type = "pdf"
            try:
                reader = PdfReader(path)
                info = reader.metadata or {}

                for k, v in info.items():
                    clean_key = str(k).lstrip("/")
                    meta[f"PDF.{clean_key}"] = str(v)

                try:
                    xmp = reader.xmp_metadata
                    if xmp is not None:
                        meta["PDF.XMP.summary"] = repr(xmp)[:2000] + "..."
                except Exception:
                    pass

                # Text scan for IPv4 in first few pages
                try:
                    text_chunks = []
                    for i, page in enumerate(reader.pages):
                        if i >= 5:
                            break
                        t = page.extract_text() or ""
                        text_chunks.append(t)
                    text = "\n".join(text_chunks)
                    ips = sorted(set(re.findall(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", text)))
                    for idx, ip in enumerate(ips, 1):
                        meta[f"TEXT.IPv4.{idx}"] = ip
                except Exception:
                    pass

            except Exception as e:
                meta["PDF.Error"] = str(e)

        # Word (.docx)
        elif suffix == ".docx":
            file_type = "word"
            try:
                doc = Document(path)
                cp = doc.core_properties

                meta["WORD.title"] = cp.title or ""
                meta["WORD.subject"] = cp.subject or ""
                meta["WORD.author"] = cp.author or ""
                meta["WORD.last_modified_by"] = cp.last_modified_by or ""
                meta["WORD.category"] = cp.category or ""
                meta["WORD.comments"] = cp.comments or ""
                meta["WORD.keywords"] = cp.keywords or ""
                meta["WORD.revision"] = str(cp.revision) if cp.revision is not None else ""
                meta["WORD.created"] = cp.created.isoformat() if cp.created else ""
                meta["WORD.modified"] = cp.modified.isoformat() if cp.modified else ""
                meta["WORD.last_printed"] = cp.last_printed.isoformat() if cp.last_printed else ""

                # Text scan for IPv4
                try:
                    text_parts = [p.text for p in doc.paragraphs]
                    text = "\n".join(text_parts)
                    ips = sorted(set(re.findall(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", text)))
                    for idx, ip in enumerate(ips, 1):
                        meta[f"TEXT.IPv4.{idx}"] = ip
                except Exception:
                    pass

            except Exception as e:
                meta["WORD.Error"] = str(e)

        return meta, exif_dict, file_type

    def render_meta(self):
        self.text.delete("1.0", tk.END)

        if not self.current_meta:
            self.text.insert(tk.END, "No metadata loaded.\n")
            return

        sections = {
            "File system metadata": [],
            "Image metadata": [],
            "EXIF metadata": [],
            "PDF metadata": [],
            "Word metadata": [],
            "Text scan (possible IPs)": [],
            "Other": [],
        }

        for k in sorted(self.current_meta.keys()):
            v = str(self.current_meta[k])
            if k.startswith("FS."):
                sections["File system metadata"].append((k, v))
            elif k.startswith("IMG."):
                sections["Image metadata"].append((k, v))
            elif k.startswith("EXIF."):
                sections["EXIF metadata"].append((k, v))
            elif k.startswith("PDF."):
                sections["PDF metadata"].append((k, v))
            elif k.startswith("WORD."):
                sections["Word metadata"].append((k, v))
            elif k.startswith("TEXT.IPv4."):
                sections["Text scan (possible IPs)"].append((k, v))
            else:
                sections["Other"].append((k, v))

        for section_title, items in sections.items():
            if not items:
                continue
            self._render_section(section_title, items)
            self.text.insert(tk.END, "\n")

        self.text.insert(
            tk.END,
            "Hints:\n"
            "  • EXIF keys: EXIF.0th.ImageDescription, EXIF.0th.Artist, EXIF.Exif.UserComment, EXIF.GPS.decimal_* etc.\n"
            "  • PDF keys: PDF.Author, PDF.Title, PDF.Creator, PDF.Producer.\n"
            "  • Word keys: WORD.author, WORD.title, WORD.last_modified_by.\n"
            "  • TEXT.IPv4.* are IP like patterns found in visible text.\n"
        )

    def _render_section(self, title, items):
        self.text.insert(tk.END, f"=== {title} ===\n")

        keys = [k for k, _ in items]
        values = [v for _, v in items]

        key_width = max(len(k) for k in keys) if keys else 0
        val_width = max(len(v) for v in values) if values else 0

        border = "+" + "-" * (key_width + 2) + "+" + "-" * (val_width + 2) + "+"

        self.text.insert(tk.END, border + "\n")
        for k, v in items:
            line = "| " + k.ljust(key_width) + " | " + v.ljust(val_width) + " |\n"
            self.text.insert(tk.END, line)
            self.text.insert(tk.END, border + "\n")

    def strip_metadata(self):
        if not self.current_path:
            messagebox.showwarning("No file", "Please load a file first.")
            return

        p = Path(self.current_path)
        clean_path = p.with_name(p.stem + "_clean" + p.suffix)

        try:
            if self.file_type == "image":
                img = Image.open(self.current_path)
                data = list(img.getdata())
                new_img = Image.new(img.mode, img.size)
                new_img.putdata(data)
                new_img.save(clean_path)
                msg = f"Clean image saved as:\n{clean_path}"

            elif self.file_type == "pdf":
                reader = PdfReader(self.current_path)
                writer = PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                writer.add_metadata({})
                with open(clean_path, "wb") as f:
                    writer.write(f)
                msg = f"Clean PDF saved as:\n{clean_path}"

            elif self.file_type == "word":
                doc = Document(self.current_path)
                cp = doc.core_properties
                cp.title = ""
                cp.subject = ""
                cp.author = ""
                cp.last_modified_by = ""
                cp.category = ""
                cp.comments = ""
                cp.keywords = ""
                doc.save(clean_path)
                msg = f"Clean Word file saved as:\n{clean_path}"

            else:
                messagebox.showinfo(
                    "Unsupported",
                    "Strip metadata is implemented for images, PDF, and .docx only."
                )
                return

            messagebox.showinfo("Success", msg)

            self.current_path = str(clean_path)
            self.file_label.config(text=str(clean_path))
            self.current_meta, self.exif_dict, self.file_type = self.read_metadata(self.current_path)
            self.pending_changes.clear()
            self.render_meta()

        except PermissionError as e:
            messagebox.showerror(
                "Permission denied",
                "The system refused write access to this file.\n\n"
                "Try this:\n"
                "  - Close any app that has the file open.\n"
                "  - Turn off Explorer preview pane.\n"
                "  - Copy the file to Documents or Desktop.\n"
                "  - Run this tool on the copy.\n\n"
                f"Details:\n{e}"
            )
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Could not strip metadata.\n\n{e}"
            )

    def edit_metadata(self):
        if not self.current_meta:
            messagebox.showwarning("No metadata", "Load a file first.")
            return

        key = simpledialog.askstring(
            "Edit metadata key",
            "Enter metadata key to edit.\n\n"
            "Examples:\n"
            "  EXIF.0th.ImageDescription\n"
            "  EXIF.0th.Artist\n"
            "  PDF.Author\n"
            "  PDF.Title\n"
            "  WORD.author\n"
            "  WORD.title\n\n"
            "You can copy a key from the table.",
        )
        if not key:
            return

        if key not in self.current_meta:
            messagebox.showerror("Key not found", "That key is not in the current metadata.")
            return

        current_value = self.current_meta.get(key, "")
        new_value = simpledialog.askstring(
            "New value",
            f"Current value for {key}:\n\n{current_value}\n\nEnter new value:",
        )
        if new_value is None:
            return

        self.current_meta[key] = new_value
        self.pending_changes[key] = new_value
        self.render_meta()

    def save_edits(self):
        if not self.current_path:
            messagebox.showwarning("No file", "Please load a file first.")
            return

        if not self.pending_changes:
            messagebox.showinfo("No changes", "There are no pending metadata edits.")
            return

        try:
            if self.file_type == "image":
                self._apply_image_edits()
            elif self.file_type == "pdf":
                self._apply_pdf_edits()
            elif self.file_type == "word":
                self._apply_word_edits()
            else:
                messagebox.showinfo("Unsupported", "Edits are supported for images, PDF, and .docx only.")
                return

            messagebox.showinfo("Saved", "Metadata edits applied.")
            self.pending_changes.clear()
            self.current_meta, self.exif_dict, self.file_type = self.read_metadata(self.current_path)
            self.render_meta()
        except Exception as e:
            messagebox.showerror("Error", f"Could not save metadata edits.\n\n{e}")

    def _apply_image_edits(self):
        if not self.exif_dict:
            messagebox.showerror(
                "Not an EXIF image",
                "There is no EXIF metadata to edit for this file.\n"
                "Edits only apply to EXIF enabled image formats.",
            )
            return

        for key, new_val in list(self.pending_changes.items()):
            if not key.startswith("EXIF.") or key.count(".") < 2:
                continue

            _, ifd_name, tag_name = key.split(".", 2)
            if ifd_name not in self.exif_dict:
                continue

            tag_id_match = None
            for tag_id, info in piexif.TAGS.get(ifd_name, {}).items():
                if info.get("name") == tag_name:
                    tag_id_match = tag_id
                    break

            if tag_id_match is None:
                continue

            try:
                self.exif_dict[ifd_name][tag_id_match] = new_val.encode("utf-8", errors="replace")
            except Exception:
                self.exif_dict[ifd_name][tag_id_match] = new_val

        exif_bytes = piexif.dump(self.exif_dict)
        img = Image.open(self.current_path)
        img.save(self.current_path, exif=exif_bytes)

    def _apply_pdf_edits(self):
        reader = PdfReader(self.current_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        metadata = dict(reader.metadata or {})
        for key, new_val in self.pending_changes.items():
            if not key.startswith("PDF.") or key.count(".") < 1:
                continue
            _, pdf_key = key.split("PDF.", 1)
            pdf_meta_key = "/" + pdf_key
            metadata[pdf_meta_key] = new_val

        writer.add_metadata(metadata)
        with open(self.current_path, "wb") as f:
            writer.write(f)

    def _apply_word_edits(self):
        doc = Document(self.current_path)
        cp = doc.core_properties

        for key, new_val in self.pending_changes.items():
            if not key.startswith("WORD."):
                continue
            field = key.split("WORD.", 1)[1]

            if field == "title":
                cp.title = new_val
            elif field == "subject":
                cp.subject = new_val
            elif field == "author":
                cp.author = new_val
            elif field == "last_modified_by":
                cp.last_modified_by = new_val
            elif field == "category":
                cp.category = new_val
            elif field == "comments":
                cp.comments = new_val
            elif field == "keywords":
                cp.keywords = new_val

        doc.save(self.current_path)


if __name__ == "__main__":
    root = tk.Tk()
    app = MetadataGUI(root)
    root.geometry("1200x650")
    root.mainloop()
